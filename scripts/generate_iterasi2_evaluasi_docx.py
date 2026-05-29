from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path


def add_heading_centered(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


def main():
    project_root = Path(__file__).resolve().parents[1]
    output_dir = project_root / "docs"
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / "Ringkasan_Evaluasi_Validasi_Iterasi_2.docx"

    doc = Document()

    # Title
    add_heading_centered(doc, "RINGKASAN EVALUASI VALIDASI ITERASI 2", level=1)
    doc.add_paragraph()

    # Content
    add_paragraph(
        doc,
        "Berdasarkan hasil percobaan pengguna, pengujian blackbox, dan UAT, Iterasi 2 menunjukkan bahwa fitur yang "
        "dikembangkan telah berjalan sesuai kebutuhan proses bisnis UKM. Fokus perbaikan Iterasi 2 meliputi: "
        "(1) batch penerimaan pendaftar diklat melalui fitur \"Terima Semua Pending\", "
        "(2) simplifikasi status anggota menjadi Active/Inactive tanpa status \"Keluar\", dan "
        "(3) penyempurnaan form pengurus dengan periode wajib serta smart filtering anggota aktif."
    )

    add_paragraph(
        doc,
        "Dari sisi UAT (percobaan user), seluruh skenario pada modul yang terkait Iterasi 2 dinyatakan berhasil, "
        "yaitu Pendaftaran Diklat (5/5), Data Anggota UKM (5/5), dan Struktur Pengurus (6/6), sehingga total validasi "
        "langsung untuk fitur inti Iterasi 2 adalah 16 skenario dengan status pass. Secara keseluruhan, sistem juga "
        "mencapai 113/113 skenario pass pada UAT lintas modul."
    )

    add_paragraph(
        doc,
        "Dari sisi blackbox testing, skenario uji fungsional untuk tiga area Iterasi 2 telah disusun dan mencakup "
        "alur input-proses-output utama, termasuk pendaftaran diklat, data anggota, dan struktur pengurus. "
        "Hal ini menunjukkan bahwa cakupan validasi fungsional telah tersedia dan selaras dengan kebutuhan pengujian sistem."
    )

    add_paragraph(
        doc,
        "Berdasarkan hasil evaluasi tersebut, sistem pada Iterasi 2 dinyatakan siap digunakan pada lingkungan operasional "
        "karena fungsi inti berjalan sesuai ekspektasi pengguna, tidak terdapat gap kebutuhan pada fitur yang direvisi, "
        "dan alur administrasi menjadi lebih efisien, sederhana, serta konsisten secara data. "
        "Dengan demikian, Iterasi 2 telah mencapai kondisi final tanpa kebutuhan revisi fitur tambahan dan dapat dilanjutkan "
        "ke tahap implementasi penuh."
    )

    doc.save(output_file)
    print(f"Dokumen berhasil dibuat: {output_file}")


if __name__ == "__main__":
    main()
